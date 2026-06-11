"""
Run: python diagnose_docx.py path/to/your/file.docx
Shows the raw XML structure so we can understand why content is missing.
"""

import sys
from collections import Counter
from docx import Document
from docx.oxml.ns import qn
from lxml import etree

def tag_local(tag):
    """Strip namespace, return just the local name."""
    return tag.split("}")[-1] if "}" in tag else tag

def walk_and_count(element, depth=0, max_depth=4, counter=None):
    if counter is None:
        counter = Counter()
    local = tag_local(element.tag)
    counter[local] += 1
    if depth < max_depth:
        for child in element:
            walk_and_count(child, depth + 1, max_depth, counter)
    return counter

def xpath_count(element, xpath_tag):
    """Count all descendants matching a tag (recursive)."""
    ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    return len(element.findall(f".//{{{ns}}}{xpath_tag}"))

def main(path):
    doc = Document(path)
    body = doc.element.body

    print(f"\n{'='*60}")
    print(f"FILE: {path}")
    print(f"{'='*60}\n")

    # 1. Direct body children
    direct_children = list(body.iterchildren())
    tag_counts = Counter(tag_local(c.tag) for c in direct_children)
    print(f"Direct <w:body> children: {len(direct_children)} total")
    for tag, count in tag_counts.most_common():
        print(f"  <w:{tag}>  x{count}")

    # 2. Recursive counts of key elements
    print(f"\nRecursive search (all descendants of <w:body>):")
    for tag in ["p", "tbl", "sdt", "sdtContent", "r", "t", "txbxContent", "altChunk"]:
        n = xpath_count(body, tag)
        if n:
            print(f"  <w:{tag}>  x{n}")

    # 3. Total text via xpath
    ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    all_t = body.findall(f".//{{{ns}}}t")
    total_text = "".join(t.text or "" for t in all_t)
    print(f"\nTotal extractable text (all <w:t> in body): {len(total_text)} chars")
    if total_text:
        preview = total_text[:300].replace("\n", " ")
        print(f"  Preview: {preview!r}")

    # 4. Headers/footers (not in body — separate parts)
    header_text = ""
    for part in doc.part.package.iter_parts():
        if "header" in part.partname or "footer" in part.partname:
            try:
                all_t2 = part.element.findall(f".//{{{ns}}}t")
                header_text += "".join(t.text or "" for t in all_t2)
            except Exception:
                pass
    print(f"\nText in headers/footers: {len(header_text)} chars")

    # 5. Text boxes (txbxContent)
    txbx_els = body.findall(f".//{{{ns}}}txbxContent")
    if txbx_els:
        txbx_text = ""
        for el in txbx_els:
            all_t3 = el.findall(f".//{{{ns}}}t")
            txbx_text += "".join(t.text or "" for t in all_t3)
        print(f"Text in text boxes (txbxContent): {len(txbx_text)} chars")
        if txbx_text:
            print(f"  Preview: {txbx_text[:200]!r}")

    # 6. doc.paragraphs count (python-docx built-in)
    print(f"\npython-docx doc.paragraphs count: {len(doc.paragraphs)}")
    non_empty = [p for p in doc.paragraphs if p.text.strip()]
    print(f"  Non-empty paragraphs: {len(non_empty)}")
    if non_empty:
        print(f"  First paragraph: {non_empty[0].text[:100]!r}")

    # 7. Show raw XML of first 3 direct body children for inspection
    print(f"\nRaw XML of first 3 direct body children:")
    for i, child in enumerate(direct_children[:3]):
        xml = etree.tostring(child, pretty_print=True).decode()
        print(f"\n  [Child {i}] <w:{tag_local(child.tag)}>")
        print("  " + "\n  ".join(xml.splitlines()[:15]))

    print(f"\n{'='*60}")
    print("DIAGNOSIS COMPLETE")
    print(f"{'='*60}\n")

if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python diagnose_docx.py <file.docx>")
        sys.exit(1)
    main(sys.argv[1])
